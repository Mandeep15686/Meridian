"""Document ingestion: extract text → semantic chunk → embed → upsert to pgvector."""

from __future__ import annotations

import hashlib
import logging
from functools import lru_cache
from pathlib import Path
from typing import Any

from openai import AsyncOpenAI
from sqlalchemy import select, text
from sqlalchemy.ext.asyncio import AsyncSession

from src.config import settings
from src.db.models import Document

logger = logging.getLogger(__name__)


@lru_cache(maxsize=1)
def get_openai_client() -> AsyncOpenAI:
    return AsyncOpenAI(api_key=settings.OPENAI_API_KEY)


# ── Text extraction ───────────────────────────────────────────────────────────


def extract_text_from_pdf(path: Path) -> tuple[str, int]:
    """Extract text from a PDF. Returns (text, page_count)."""
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages), len(reader.pages)


def extract_text_from_docx(path: Path) -> str:
    """Extract text from a DOCX file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text(path: Path, mime_type: str = "") -> tuple[str, int]:
    """
    Dispatch to the correct extractor based on MIME type or extension.

    Returns:
        (text_content, page_count) — page_count is 0 for non-page formats.
    """
    ext = path.suffix.lower()
    if ext == ".pdf" or "pdf" in mime_type:
        return extract_text_from_pdf(path)
    if ext in (".docx", ".doc") or "word" in mime_type:
        return extract_text_from_docx(path), 0
    if ext in (".txt", ".md", ".html"):
        return path.read_text(encoding="utf-8", errors="replace"), 0
    logger.warning("No specific extractor for %s — reading as plain text", path.name)
    return path.read_text(encoding="utf-8", errors="replace"), 0


# ── Semantic chunking ─────────────────────────────────────────────────────────


def semantic_chunk(
    text: str,
    target_size: int | None = None,
    overlap: int | None = None,
) -> list[str]:
    """
    Split text into semantically coherent chunks.

    Uses LlamaIndex SemanticChunker when available; falls back to
    sentence-boundary splitting.
    """
    target_size = target_size or settings.CHUNK_SIZE_TARGET
    overlap = overlap or settings.CHUNK_OVERLAP

    try:
        from llama_index.core import Document as LlamaDoc
        from llama_index.core.node_parser import SemanticSplitterNodeParser
        from llama_index.embeddings.openai import OpenAIEmbedding

        embed_model = OpenAIEmbedding(model=settings.EMBEDDING_MODEL)
        splitter = SemanticSplitterNodeParser(
            buffer_size=1,
            breakpoint_percentile_threshold=95,
            embed_model=embed_model,
        )
        doc = LlamaDoc(text=text)
        nodes = splitter.get_nodes_from_documents([doc])
        return [n.get_content() for n in nodes if n.get_content().strip()]

    except Exception as exc:
        logger.warning("SemanticChunker unavailable (%s); using sentence splitter", exc)
        return _sentence_chunk(text, target_size, overlap)


def _sentence_chunk(text: str, target_size: int, overlap: int) -> list[str]:
    """Fallback: split on sentence boundaries targeting token count."""
    import re

    sentences = re.split(r"(?<=[.!?])\s+", text)
    chunks: list[str] = []
    current: list[str] = []
    current_len = 0
    chars_per_token = 4

    for sentence in sentences:
        sentence_len = len(sentence) // chars_per_token
        if current_len + sentence_len > target_size and current:
            chunks.append(" ".join(current))
            # Keep overlap sentences at start of next chunk
            overlap_sentences = current[-(overlap // 20 or 1) :]
            current = overlap_sentences + [sentence]
            current_len = sum(len(s) // chars_per_token for s in current)
        else:
            current.append(sentence)
            current_len += sentence_len

    if current:
        chunks.append(" ".join(current))

    return [c.strip() for c in chunks if c.strip()]


# ── Embedding ────────────────────────────────────────────────────────────────


async def embed_texts(texts: list[str]) -> list[list[float]]:
    """Batch embed texts using the configured embedding model."""
    if not texts:
        return []

    # OpenAI embedding API accepts up to 2048 inputs per request
    batch_size = 512
    all_embeddings: list[list[float]] = []

    for i in range(0, len(texts), batch_size):
        batch = [t[:8191] for t in texts[i : i + batch_size]]
        response = await get_openai_client().embeddings.create(
            model=settings.EMBEDDING_MODEL,
            input=batch,
        )
        all_embeddings.extend([item.embedding for item in response.data])

    return all_embeddings


# ── Database upsert ───────────────────────────────────────────────────────────


async def upsert_chunks(
    session: AsyncSession,
    corpus_id: str,
    document_id: str,
    regulation: str,
    jurisdiction: str,
    chunks: list[str],
    embeddings: list[list[float]],
    article_metadata: list[dict[str, Any]] | None = None,
) -> int:
    """
    Upsert chunk rows into the chunks table.

    Uses ON CONFLICT DO NOTHING for idempotent re-ingestion.
    Returns number of chunks actually inserted.
    """
    if not chunks or len(chunks) != len(embeddings):
        raise ValueError("chunks and embeddings must be the same length")

    metadata = article_metadata or [{}] * len(chunks)
    inserted = 0

    for i, (content, embedding, meta) in enumerate(zip(chunks, embeddings, metadata, strict=True)):
        vec_str = "[" + ",".join(str(v) for v in embedding) + "]"

        stmt = text(
            """
            INSERT INTO chunks (
                id, document_id, corpus_id, regulation, article, article_title,
                jurisdiction, chunk_index, content, token_count, embedding
            )
            VALUES (
                gen_random_uuid(), :document_id, :corpus_id, :regulation,
                :article, :article_title, :jurisdiction, :chunk_index,
                :content, :token_count, :embedding::vector
            )
            ON CONFLICT (document_id, chunk_index) DO NOTHING
            RETURNING id;
        """
        )

        result = await session.execute(
            stmt,
            {
                "document_id": document_id,
                "corpus_id": corpus_id,
                "regulation": regulation,
                "article": meta.get("article"),
                "article_title": meta.get("article_title"),
                "jurisdiction": jurisdiction,
                "chunk_index": i,
                "content": content,
                "token_count": len(content) // 4,
                "embedding": vec_str,
            },
        )
        if result.scalar_one_or_none() is not None:
            inserted += 1

    return inserted


# ── Full ingestion pipeline ───────────────────────────────────────────────────


async def ingest_document(
    session: AsyncSession,
    corpus_id: str,
    regulation: str,
    jurisdiction: str,
    source_path: Path,
    source_url: str | None = None,
    article_metadata_fn: Any | None = None,
    force_reingest: bool = False,
) -> dict[str, Any]:
    """
    Full ingestion pipeline for a single regulatory document.

    Args:
        session: Async DB session.
        corpus_id: ID of the parent Corpus row.
        regulation: Short corpus slug, e.g. 'gdpr'.
        jurisdiction: 'EU', 'US', 'global', etc.
        source_path: Local path to the document file.
        source_url: Canonical URL for citation.
        article_metadata_fn: Optional callable(chunk_text) → {article, article_title}.
        force_reingest: Re-ingest even if content_hash already exists.

    Returns:
        Dict with keys: document_id, chunks_inserted, chunks_total.
    """
    # Hash the document content for deduplication
    content_hash = hashlib.sha256(source_path.read_bytes()).hexdigest()

    # Check for existing document
    existing = await session.execute(
        select(Document).where(
            Document.corpus_id == corpus_id,
            Document.content_hash == content_hash,
        )
    )
    existing_doc = existing.scalar_one_or_none()

    if existing_doc and not force_reingest:
        logger.info("Document %s already ingested (hash match), skipping", source_path.name)
        return {
            "document_id": existing_doc.id,
            "chunks_inserted": 0,
            "chunks_total": 0,
            "skipped": True,
        }

    # Extract text
    text_content, page_count = extract_text(source_path)
    if not text_content.strip():
        logger.warning("Empty text extracted from %s", source_path.name)
        return {
            "document_id": None,
            "chunks_inserted": 0,
            "chunks_total": 0,
            "skipped": True,
        }

    token_count = len(text_content) // 4

    # Create document record
    doc = Document(
        corpus_id=corpus_id,
        filename=source_path.name,
        source_url=source_url,
        content_hash=content_hash,
        page_count=page_count or None,
        token_count=token_count,
    )
    session.add(doc)
    await session.flush()  # get doc.id

    # Chunk text
    logger.info("Chunking %s (%d tokens)...", source_path.name, token_count)
    chunks = semantic_chunk(text_content)
    logger.info("Produced %d chunks from %s", len(chunks), source_path.name)

    # Build article metadata per chunk
    article_metas: list[dict[str, Any]] = []
    for chunk in chunks:
        if article_metadata_fn:
            article_metas.append(article_metadata_fn(chunk))
        else:
            article_metas.append({})

    # Embed chunks
    logger.info("Embedding %d chunks...", len(chunks))
    embeddings = await embed_texts(chunks)

    # Upsert into DB
    inserted = await upsert_chunks(
        session=session,
        corpus_id=corpus_id,
        document_id=doc.id,
        regulation=regulation,
        jurisdiction=jurisdiction,
        chunks=chunks,
        embeddings=embeddings,
        article_metadata=article_metas,
    )

    await session.commit()

    logger.info("Ingested %s: %d/%d chunks inserted", source_path.name, inserted, len(chunks))
    return {
        "document_id": doc.id,
        "chunks_inserted": inserted,
        "chunks_total": len(chunks),
        "skipped": False,
    }
